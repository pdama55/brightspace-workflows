"""Turn a syllabus into a date -> topic table.

Brightspace has no structured record of what a class is *doing* on a given day —
that lives in a syllabus PDF a human wrote. So we extract it once per term with
Claude and cache the result; the daily digest just looks today's date up.
"""
from __future__ import annotations

import base64
import json
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional

from .config import SCHEDULE_DIR

SCHEMA = {
    "type": "object",
    "properties": {
        "course": {"type": "string"},
        "meetings": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "date": {
                        "type": ["string", "null"],
                        "description": "ISO date (YYYY-MM-DD), or null if the syllabus gives no resolvable date.",
                    },
                    "label": {
                        "type": ["string", "null"],
                        "description": "How the syllabus refers to this session, e.g. 'Week 3, Tuesday'.",
                    },
                    "topic": {"type": "string"},
                    "readings": {"type": ["string", "null"]},
                    "notes": {
                        "type": ["string", "null"],
                        "description": "Anything else due or expected that day. Null if nothing.",
                    },
                },
                "required": ["date", "label", "topic", "readings", "notes"],
                "additionalProperties": False,
            },
        },
    },
    "required": ["course", "meetings"],
    "additionalProperties": False,
}

PROMPT = """Extract the class-by-class schedule from this syllabus.

Course label to use: {course}
{term_hint}
Rules:
- One entry per class meeting, in chronological order.
- `date` must be an ISO date (YYYY-MM-DD). If the syllabus only gives relative
  positions ("Week 3, Tuesday") and you were given a term start date above,
  resolve it to a real date. If you cannot resolve it, set `date` to null and put
  the syllabus's own wording in `label`.
- `topic` is what the class covers that day, in the syllabus's own words, trimmed.
- `readings` is assigned reading/viewing for that day, or null.
- `notes` is anything else due or expected that day, or null.
- Do not invent meetings that the syllabus does not list. If the syllabus has no
  class-by-class schedule at all, return an empty `meetings` array."""


@dataclass
class Meeting:
    date: Optional[str]
    label: Optional[str]
    topic: str
    readings: Optional[str]
    notes: Optional[str]

    def render(self) -> str:
        parts = [self.topic.strip()]
        if self.readings:
            parts.append(f"read {self.readings.strip()}")
        if self.notes:
            parts.append(self.notes.strip())
        return " · ".join(p for p in parts if p)


def slugify(name: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", name.strip().lower()).strip("-")
    return slug or "course"


def schedule_path(course: str) -> Path:
    return SCHEDULE_DIR / f"{slugify(course)}.json"


def _document_block(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        encoded = base64.standard_b64encode(path.read_bytes()).decode("ascii")
        return {
            "type": "document",
            "source": {"type": "base64", "media_type": "application/pdf", "data": encoded},
        }
    if suffix == ".docx":
        # Half of Brightspace syllabi are .docx; extract the text rather than
        # making the user convert each one by hand.
        return {
            "type": "document",
            "source": {"type": "text", "media_type": "text/plain", "data": _docx_text(path)},
        }
    if suffix in {".txt", ".md", ".markdown", ".rst"}:
        return {
            "type": "document",
            "source": {
                "type": "text",
                "media_type": "text/plain",
                "data": path.read_text(encoding="utf-8", errors="replace"),
            },
        }
    raise ValueError(
        f"Unsupported syllabus format {suffix!r}. Use a .pdf, .docx, or plain-text file."
    )


def _docx_text(path: Path) -> str:
    """Paragraphs plus table cells — syllabus schedules are almost always tables."""
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "Reading .docx needs python-docx. Run: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    chunks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))
    text = "\n".join(chunks)
    if not text.strip():
        raise RuntimeError(f"{path.name} produced no extractable text.")
    return text


def parse_syllabus(
    path: Path,
    course: str,
    model: str,
    term_start: Optional[date] = None,
) -> dict:
    """Ask Claude for the schedule table. Returns the raw parsed dict."""
    import os

    import anthropic

    # The SDK defers this error until request time, which buries it under a
    # streaming traceback. Check up front so the fix is obvious.
    if not (os.getenv("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_AUTH_TOKEN")):
        raise RuntimeError(
            "ANTHROPIC_API_KEY is not set. Add it to your .env — get a key at "
            "https://console.anthropic.com/settings/keys. This is only needed for "
            "parse-syllabus; the daily digest does not use it."
        )
    client = anthropic.Anthropic()

    term_hint = (
        f"The term starts on {term_start.isoformat()} (a {term_start.strftime('%A')}). "
        "Use it to resolve week numbers to real dates.\n"
        if term_start
        else "No term start date was provided.\n"
    )

    with client.messages.stream(
        model=model,
        max_tokens=32000,
        output_config={"effort": "medium", "format": {"type": "json_schema", "schema": SCHEMA}},
        messages=[
            {
                "role": "user",
                "content": [
                    _document_block(path),
                    {"type": "text", "text": PROMPT.format(course=course, term_hint=term_hint)},
                ],
            }
        ],
    ) as stream:
        response = stream.get_final_message()

    if response.stop_reason == "max_tokens":
        raise RuntimeError(
            "Claude hit the output limit before finishing the schedule. The syllabus is "
            "unusually long — split it and run parse-syllabus once per part."
        )

    text = next((b.text for b in response.content if b.type == "text"), None)
    if not text:
        raise RuntimeError(f"No JSON came back from the model (stop_reason={response.stop_reason}).")
    return json.loads(text)


def save_schedule(course: str, data: dict) -> Path:
    SCHEDULE_DIR.mkdir(parents=True, exist_ok=True)
    path = schedule_path(course)
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")
    return path


def load_schedules() -> Dict[str, List[Meeting]]:
    """Load every cached schedule, keyed by course label (lowercased)."""
    schedules: Dict[str, List[Meeting]] = {}
    if not SCHEDULE_DIR.exists():
        return schedules
    for path in sorted(SCHEDULE_DIR.glob("*.json")):
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            continue
        course = str(data.get("course") or path.stem)
        meetings = [
            Meeting(
                date=m.get("date"),
                label=m.get("label"),
                topic=str(m.get("topic") or "").strip(),
                readings=m.get("readings"),
                notes=m.get("notes"),
            )
            for m in data.get("meetings", [])
            if str(m.get("topic") or "").strip()
        ]
        schedules[course.lower()] = meetings
    return schedules


def topics_on(schedules: Dict[str, List[Meeting]], day: date) -> List[tuple]:
    """Every course's plan for a given day, as (course, text).

    Used when the calendar feed carries no class-meeting events — which is the
    norm on Brightspace, since most instances only publish due dates.
    """
    target = day.isoformat()
    out: List[tuple] = []
    for course, meetings in sorted(schedules.items()):
        for meeting in meetings:
            if meeting.date == target:
                # Schedules are keyed lowercase for matching; course codes display
                # uppercase everywhere else, so normalize on the way out.
                out.append((course.upper(), meeting.render()))
                break
    return out


def topic_for(schedules: Dict[str, List[Meeting]], course: Optional[str], day: date) -> Optional[str]:
    """Find today's topic for a course, matching course names loosely."""
    if not course:
        return None
    needle = course.strip().lower()
    target = day.isoformat()

    for key, meetings in schedules.items():
        if key != needle and key not in needle and needle not in key:
            continue
        for meeting in meetings:
            if meeting.date == target:
                return meeting.render()
    return None
